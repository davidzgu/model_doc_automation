import os
import datetime
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from typing import Optional, List, Dict
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from langchain_core.messages import HumanMessage, SystemMessage

from bsm_multi_agents.graph.state import WorkflowState
from bsm_multi_agents.config.llm_config import get_llm
from bsm_multi_agents.agents.utils import invoke_llm

def report_generator_node(state: WorkflowState) -> WorkflowState:
    """
    LangGraph node: Generates the final Word report directly from validated CSV data.
    
    1. Reads the CSV (input or validated results).
    2. Calculates statistical summaries (Token Optimization).
    3. Uses LLM (`get_llm`) to generate Section 2 (Summary) and Section 3 (Pricing Analysis).
    4. Generates charts and tables directly.
    5. Saves 'final_OPA_report.docx' to output_dir.
    """
    errors = state.get("errors", [])
    
    # --- 1. Input Resolution ---
    # Prioritize validated results if available, else standard input
    csv_path = state.get("validate_results_path")
    if not csv_path or not os.path.exists(csv_path):
        # Fallback to original input if validate_results (or greeks_results) missing
        # The user mentioned 'dummy_options_greeks_results_validate_results.csv' specifically.
        # We try strict fallback chain: validate -> greeks -> input
        csv_path = state.get("greeks_results_path")
        if not csv_path or not os.path.exists(csv_path):
            csv_path = state.get("csv_file_path")
            
    output_dir = state.get("output_dir")

    if not csv_path or not os.path.exists(csv_path):
        errors.append("report_generator_node: No valid CSV input found (checked validate/greeks/input paths)")
        state["errors"] = errors
        return state
        
    if not output_dir:
        errors.append("report_generator_node: output_dir missing")
        state["errors"] = errors
        return state

    try:
        # --- 2. Data Loading & Stats (Token Opt) ---
        df = pd.read_csv(csv_path)
        stats_summary = _calculate_stats_summary(df)
        
        # --- 3. Content Generation (LLM) ---
        llm = get_llm()
        
        # Section 2: Executive Summary (based on overall stats)
        section2_content = _generate_section2(llm, stats_summary)
        
        # Section 3: Pricing Analysis (Iterative per asset class)
        # We'll generate this inside the doc building loop or pre-generate. 
        # Let's pre-generate to separate logic.
        asset_analyses = {}
        assets = df["asset_class"].unique() if "asset_class" in df.columns else []
        for asset in assets:
            df_asset = df[df["asset_class"] == asset]
            asset_stats = _calculate_stats_summary(df_asset)
            asset_analyses[asset] = _generate_asset_analysis(llm, asset, asset_stats)

        # --- 4. Document Building ---
        output_filename = "final_OPA_report.docx"
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, output_filename)
        
        doc = Document()
        _add_title_page(doc)
        _add_table_of_contents(doc)
        _add_section1(doc)
        
        # Section 2
        _add_section_content(doc, "2. Summary of Analysis", section2_content)
        
        # Section 3
        doc.add_heading("3. Pricing Analysis", level=1)
        if not assets.size:
             doc.add_paragraph("No asset class data found in CSV.")
        else:
            for asset in assets:
                _add_asset_section(doc, df, asset, asset_analyses.get(asset, ""))

        # Save
        doc.save(output_path)
        state["final_report_path"] = os.path.abspath(output_path)

    except Exception as e:
        errors.append(f"report_generator_node: Error execution: {e}")

    state["errors"] = errors
    return state


# --- Internal Helpers: Data & LLM ---

def _calculate_stats_summary(df: pd.DataFrame) -> str:
    """
    Generates a textual statistical summary of the DataFrame to save tokens.
    Avoids passing raw rows.
    """
    if df.empty:
        return "Dataset is empty."
        
    stats = []
    stats.append(f"Total Rows: {len(df)}")
    
    # Columns presence
    stats.append(f"Columns: {', '.join(df.columns)}")
    
    # Numerical summaries
    num_cols = df.select_dtypes(include=['float', 'int']).columns
    if not num_cols.empty:
        desc = df[num_cols].describe().to_string()
        stats.append(f"Numerical Stats:\n{desc}")
    
    # Subsets (Calls vs Puts count)
    if "option_type" in df.columns:
        counts = df["option_type"].value_counts().to_string()
        stats.append(f"Option Type Distribution:\n{counts}")
        
    return "\n\n".join(stats)



def _generate_section2(llm, stats_summary: str) -> str:
    system_prompt = (
        "You are a quantitative finance expert. "
        "Write an Executive Summary ('Section 2') for an Option Pricing Analysis Report. "
        "Base it on the provided statistical summary of the results. "
        "Highlight data scale, key average values (like average price), and distribution coverage. "
        "Keep it professional, concise, and narrative."
    )
    user_prompt = f"Data Statistics:\n{stats_summary}"
    messages = [
        SystemMessage(content=system_prompt),
        HumanMessage(content=user_prompt)
    ]
    response = llm.invoke(messages)
    return response

def _generate_asset_analysis(llm, asset: str, stats_summary: str) -> str:
    system_prompt = (
        "You are a quantitative analyst. Provide a specific analysis for the given Asset Class. "
        "Discuss the statistical profile provided (e.g. average prices, volatility ranges). "
        "Mention if the data covers a wide range of maturities (T) or Strikes (K) based on the stats."
    )
    user_prompt = f"Asset Class: {asset}\n\nStats:\n{stats_summary}"
    messages = [
        SystemMessage(content=system_prompt),
        HumanMessage(content=user_prompt)
    ]
    response = llm.invoke(messages)
    return response


# --- Internal Helpers: Document Building ---

def _add_title_page(doc):
    p = doc.add_paragraph()
    run = p.add_run("Ongoing Monitoring Analysis Report")
    run.bold = True
    run.font.size = doc.styles["Title"].font.size
    p.alignment = 1

    p = doc.add_paragraph()
    p.add_run("Option Pricing, BSM").bold = True
    p.alignment = 1
    
    doc.add_paragraph("")
    
    today = datetime.date.today().strftime("%B %d, %Y")
    for l, v in [("Author", "John Doe"), ("Group", "Quant Analytics"), ("Date", today)]:
        p = doc.add_paragraph()
        p.add_run(f"{l}: ").bold = True
        p.add_run(v)
        p.alignment = 1
    doc.add_page_break()

def _add_table_of_contents(doc):
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    p._p.append(fld)
    doc.add_page_break()

def _add_section1(doc):
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This report summarizes the validation and pricing results of the BSM model. "
        "The analysis is based on the input dataset and validating the calculated Greeks."
    )

def _add_section_content(doc, title, content):
    doc.add_heading(title, level=1)
    for block in content.split("\n\n"):
        if block.strip():
            doc.add_paragraph(block.strip())

def _add_asset_section(doc, df_full, asset, analysis_text):
    doc.add_heading(f"Pricing Output: {asset}", level=2)
    
    # 1. Narrative
    if analysis_text.strip():
        doc.add_paragraph(analysis_text.strip())
    
    # 2. Plot
    df_asset = df_full[df_full["asset_class"] == asset].copy()
    if df_asset.empty:
        return

    if "T" in df_asset.columns and "BSM_Price" in df_asset.columns:
        _add_plot(doc, df_asset, asset)

    # 3. Table (Limit rows for display if needed? User didn't specify, but safer to limit for large docs)
    # We'll print all rows as requested in previous iterations, assuming dataset isn't huge.
    _add_datatable(doc, df_asset)

def _add_plot(doc, df, asset):
    df = df.sort_values("T")
    fig, ax = plt.subplots(figsize=(6, 4))
    
    if "option_type" in df.columns:
        for otype in ["call", "put"]:
            subset = df[df["option_type"] == otype]
            if not subset.empty:
                ax.plot(subset["T"], subset["BSM_Price"], label=otype, marker='o')
        ax.legend()
    else:
        ax.plot(df["T"], df["BSM_Price"], marker='o')
        
    ax.set_xlabel("Maturity (T)")
    ax.set_ylabel("BSM Price")
    ax.set_title(f"{asset} Pricing Curve")
    ax.grid(True, linestyle='--', alpha=0.6)

    img_stream = BytesIO()
    fig.savefig(img_stream, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    img_stream.seek(0)
    doc.add_picture(img_stream, width=Inches(6.0))

def _add_datatable(doc, df):
    # Sanitize for table
    df_disp = df.fillna("")
    
    table = doc.add_table(rows=len(df_disp) + 1, cols=len(df_disp.columns))
    table.style = "Table Grid"
    
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df_disp.columns):
        hdr_cells[i].text = str(col)
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        
    for i, row in enumerate(df_disp.itertuples(index=False)):
        cells = table.rows[i+1].cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
