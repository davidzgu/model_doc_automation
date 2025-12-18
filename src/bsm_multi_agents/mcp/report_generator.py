import os
from typing import Optional
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openai import OpenAI
from bsm_multi_agents.config.llm_config import DEFAULT_OPENAI_API_KEY
import datetime 
import csv
import pandas as pd
from io import BytesIO
import matplotlib.pyplot as plt
from docx.shared import Inches

def summary_to_word(
    summary_docx_path: str,
    output_docx_path: str = "final_OPA_report.docx",
    pricing_path: str = "./",
    title: str = "Ongoing Monitoring Analysis Report",
    model_name: str = "Option Pricing, BSM",
    author_name: str = "John Doe",
    group_name: str = "Front Desk Modeling and Analytics",
    version: str = "v1.0",
    section1_heading: str = "1. Introduction",
    section1_paragraph: Optional[str] = None,
    section2_heading: str = "2. Testing",
    section2_paragraph: Optional[str] = None,
) -> str:
    """
    Create a formatted Word report with:
    - Title page
    - Table of contents (Word-updatable)
    - Section 1: static intro paragraph
    - Section 2: refined summary text from an existing Word file via OpenAI

    Args:
        summary_docx_path: Path to the existing .docx file containing the raw summary text.
        output_docx_path: Path for the generated formatted report (.docx).
        title: Title used on the title page and as document main title.
        section1_heading: Heading text for Section 1 (level-1 heading).
        section1_paragraph: Optional custom intro paragraph for Section 1.
                            If omitted, a default intro is used.
        section2_heading: Heading text for Section 2 (level-1 heading).
        section2_paragraph: Optional testing analysis paragraph for Section 2.

    Returns:
        Absolute path to the generated Word document.
    """

    # --------- 0. Basic checks ---------
    if not os.path.exists(summary_docx_path):
        raise FileNotFoundError(f"Summary file not found: {summary_docx_path}")

    api_key = os.getenv("OPENAI_API_KEY", DEFAULT_OPENAI_API_KEY)
    if not api_key:
        raise RuntimeError(
            "OPENAI_API_KEY is not set in the environment. "
            "Please export it before calling this tool."
        )

    # --------- 1. Load raw summary text from the source Word file ---------
    src_doc = Document(summary_docx_path)
    raw_summary_parts = []
    for p in src_doc.paragraphs:
        text = p.text.strip()
        if text:
            raw_summary_parts.append(text)
    raw_summary_text = "\n\n".join(raw_summary_parts).strip()

    if not raw_summary_text:
        raise ValueError("No non-empty text found in the summary document.")

    # --------- 2. Refine summary via OpenAI API ---------
    client = OpenAI(api_key=api_key)

    system_prompt = (
        "You are a quantitative finance expert and technical writer. "
        "Given a rough summary about a quantitative analysis, rewrite it into a clear, "
        "well-structured Section 2 of a professional report. "
        "Use concise English, with logical flow and numbered or bulleted lists where helpful. "
        "Do NOT include a title page or table of contents; only the Section 2 narrative itself."
    )

    user_prompt = (
        "Here is the raw summary text that should become Section 2 of the report. "
        "Please refine it as described:\n\n"
        f"{raw_summary_text}"
    )

    # Using chat.completions; you can swap to Responses API if you prefer
    completion = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.3,
    )
    refined_summary = completion.choices[0].message.content.strip()

    # --------- 3. Create new Word document ---------
    doc = Document()

    # Title
    title_run = doc.add_paragraph().add_run(title)
    title_run.bold = True
    title_run.font.size = doc.styles["Title"].font.size
    doc.paragraphs[-1].alignment = 1  # 0=left, 1=center

    # Subtitle (Model Name)
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(model_name)
    subtitle_para.alignment = 1
    subtitle_run.bold = True

    doc.add_paragraph("")  # spacer

    # Author
    author_para = doc.add_paragraph()
    author_para.add_run("Author: ").bold = True
    author_para.add_run(author_name).italic = False
    author_para.alignment = 1

    # Group
    group_para = doc.add_paragraph()
    group_para.add_run("Group: ").bold = True
    group_para.add_run(group_name)
    group_para.alignment = 1

    # Report date
    date_str = datetime.date.today().strftime("%B %d, %Y")
    date_para = doc.add_paragraph()
    date_para.add_run("Report Date: ").bold = True
    date_para.add_run(date_str)
    date_para.alignment = 1

    # Version
    version_para = doc.add_paragraph()
    version_para.add_run("Document Version: ").bold = True
    version_para.add_run(version)
    version_para.alignment = 1

    # Page break after title page
    doc.add_page_break()

    # ============================
    # TABLE OF CONTENTS (Three Levels)
    # ============================

    doc.add_heading("Table of Contents", level=1)

    # Add a TOC that includes Heading 1, Heading 2, and Heading 3
    toc_paragraph = doc.add_paragraph()

    # Create the TOC XML field
    toc_field = OxmlElement("w:fldSimple")
    toc_field.set(
        qn("w:instr"),
        'TOC \\o "1-3" \\h \\z \\u'  
        # \o "1-3" → include Heading levels 1 to 3
        # \h       → hyperlink each entry
        # \z       → hide page numbers in web preview (but visible in Word)
        # \u       → use applied outline levels
    )

    # Insert the TOC field into the paragraph
    toc_paragraph._p.append(toc_field)

    # Page break after TOC
    doc.add_page_break()
    
    # ============================
    # SECTION 1
    # ============================
    if not section1_paragraph:
        section1_paragraph = (
            "This section provides contextual background, objectives, and relevant "
            "considerations for the ongoing monitoring analysis. Subsequent sections "
            "expand on methodology, insights, and results."
        )

    doc.add_heading(section1_heading, level=1)
    doc.add_paragraph(section1_paragraph)

    # ============================
    # SECTION 2 – Refined Summary
    # ============================
    doc.add_heading("2. Summary of Analysis", level=1)

    

    # Read CSV
    # with open(pricing_path, newline="", encoding="utf-8") as f:
    #     reader = csv.reader(f)
    #     rows = list(reader)
    df_pricing = pd.read_csv(pricing_path)
    section_ordering = 0
    for asset in ["FX", "Equity", "Commodity"]:
        section_ordering+=1
        doc.add_heading("2." + str(section_ordering) + " Summary of Analysis for " + asset, level=2)
        doc.add_heading("2." + str(section_ordering) + ".1 Summary of Pricing for " + asset, level=3)
        doc.add_paragraph("The pricing output of "+asset+ " listed in the below table,")
        df = df_pricing[df_pricing["asset_class"] == asset]
        df = df.sort_values("T")
        df = df.dropna()
        
        fig, ax = plt.subplots()
        df_call = df[df["option_type"] == "call"]
        ax.plot(df_call["T"], df_call["BSM_Price"], label = "call")
        df_put = df[df["option_type"] == "put"]
        ax.plot(df_put["T"], df_put["BSM_Price"], label = "put")

        ax.set_xlabel("Time to Maturity (T)")
        ax.set_ylabel("Option Price (BSM)")
        ax.set_title(f"Option Pricing Curve – {asset}")  
        ax.legend()

        # Save figure to memory (PNG bytes)
        img_stream = BytesIO()
        fig.savefig(img_stream, format="png", dpi=200, bbox_inches="tight")
        plt.close(fig)
        img_stream.seek(0)

        system_prompt = (
            "Please summmarize the option pricing results (pull and call) from the tables with the following topics"
            "1. Overall Data Quality"
            "2. Pricing Level by Asset Class"
            "3. Term Structure (Price vs Maturity)"
            "4. Call vs Put Behavior"
            "5. Model Consistency"
            "6. Key Takeaway"
            "Also, we have some annotation for you to understand table columns"
            "- **Valuation Date:** The date on which the option price is calculated."
            "- **Spot Price (S):** Current price of the underlying asset."
            "- **Strike Price (K):** Exercise price of the option."
            "- **Time to Maturity (T):** Time remaining until option expiration, expressed in years."
            "- **Risk-Free Rate (r):** Annualized risk-free interest rate, used for discounting."
            "- **Volatility (σ):** Annualized standard deviation of the underlying asset’s returns."
            "- **Option Type:** Call or Put."
            "- **Asset Class:** Classification of the underlying asset (e.g., equity, index)."
        )

        user_prompt = (
            "Here is the raw pull and call price tables that should become the option pricing output summary section. The title should have this asset class name."
            "Please refine it as described:\n\n"
            f"{df_call}"
            f"{df_put}"
        )

        # Using chat.completions; you can swap to Responses API if you prefer
        completion = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.3,
        )
        refined_summary1 = completion.choices[0].message.content.strip() 
        for block in refined_summary1.split("\n\n"):
            block = block.strip()
            if block:
                doc.add_paragraph(block)        



        # Insert figure
        doc.add_heading("2." + str(section_ordering) + ".2 Visualization of Pricing for " + asset, level=3)
        doc.add_picture(img_stream, width=Inches(6.5))

        doc.add_heading("2." + str(section_ordering) + ".3 Summary of Gamma Positivity Testing for " + asset, level=3)
        system_prompt = (
            "Please summmarize the gamma positivity results from the tables with only summarizing quantitative statistics and quantitative explanation into two paragraphs."
            "Also, we have some annotation for you to understand table columns"
            "Parameters:"
            "option_type (str): 'call' for call option, 'put' for put option"
            "S (float): current stock price"
            "K (float): option strike price"
            "T (float): time to expiration in years"
            "r (float): risk-free interest rate (annualized)"
            "sigma (float): volatility of the underlying stock (annualized)"
            "gamma_positive: True if gamma positivity holds, False otherwise"
        )

        user_prompt = (
            "Here is the raw testing result tables that should become the option pricing output gamma positivity testing summary section."
            "The title should have this asset class name."
            "Please refine it as described:\n\n"
            f"{df}"
        )

        # Using chat.completions; you can swap to Responses API if you prefer
        completion = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.3,
        )
        refined_summary1 = completion.choices[0].message.content.strip() 
        for block in refined_summary1.split("\n\n"):
            block = block.strip()
            if block:
                doc.add_paragraph(block)        
       

        # # Create table (rows = data + header)
        # table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        # table.style = "Table Grid"

        # # Header row
        # for col_idx, col_name in enumerate(df.columns):
        #     cell = table.rows[0].cells[col_idx]
        #     cell.text = str(col_name)
        #     cell.paragraphs[0].runs[0].bold = True

        # # Data rows
        # for row_idx in range(df.shape[0]):
        #     for col_idx in range(df.shape[1]):
        #         table.rows[row_idx + 1].cells[col_idx].text = str(df.iat[row_idx, col_idx])


    # Paragraph AFTER table
    # doc.add_paragraph("")    

    for block in refined_summary.split("\n\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block)

    # --------- 4. Save ---------
    doc.save(output_docx_path)
    return os.path.abspath(output_docx_path)