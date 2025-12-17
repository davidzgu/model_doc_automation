from docx import Document
import csv
from typing import Optional
import os

def csv_to_summary(
    csv_path: str,
    output_dir: str,
    max_rows: int = 20
) -> str:
    """
    Read a CSV of greeks results and generate a Word document with a data table and explanation.
    """
    if not os.path.exists(csv_path):
        raise FileNotFoundError(f"CSV file not found: {csv_path}")

    with open(csv_path, newline="", encoding="utf-8") as f:
        reader = csv.reader(f)
        rows = list(reader)

    if not rows:
        raise ValueError("CSV file is empty")

    header = rows[0]
    data_rows = rows[1:max_rows + 1]

    if not data_rows:
        raise ValueError("CSV has header but no data rows")

    doc = Document()
    doc.add_heading("Black–Scholes–Merton Pricing Results", level=1)

    intro = (
        "This document summarizes sample Black–Scholes–Merton (BSM) option pricing "
        "results. The table below contains a subset of rows from the input dataset."
    )
    doc.add_paragraph(intro)

    num_rows = len(data_rows) + 1
    num_cols = len(header)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for j, col_name in enumerate(header):
        hdr_cells[j].text = str(col_name)

    for i, row in enumerate(data_rows, start=1):
        row_cells = table.rows[i].cells
        for j, value in enumerate(row):
            row_cells[j].text = "" if value is None else str(value)

    explanation = (
        "The columns include the valuation date, spot price S, strike K, time to "
        "maturity T, risk-free rate r, volatility sigma, option type, asset class, "
        "and the resulting BSM_Price. Rows with missing or invalid inputs may yield "
        "missing or unreliable prices and should be handled carefully in a "
        "production workflow."
    )
    doc.add_paragraph(explanation)

    filename = os.path.basename(csv_path).replace(".csv", "_summary.docx")
    output_path = os.path.join(output_dir, filename)

    doc.save(output_path)
    return os.path.abspath(output_path)

