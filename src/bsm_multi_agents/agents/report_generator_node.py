from typing import Optional, Dict, List
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime 
from io import BytesIO
import matplotlib.pyplot as plt
from docx.shared import Inches
from pathlib import Path
import seaborn as sns
import re
import io

from langchain_core.messages import SystemMessage, HumanMessage
from bsm_multi_agents.config.llm_config import get_llm
from bsm_multi_agents.graph.state import WorkflowState


def report_generator_agent_node(state: WorkflowState) -> WorkflowState:
    """Refactored report generator agent node"""
    print("\n>>> [Report Generator Agent] Compiling final report...")

    ## Load parameters
    errors = state.get("errors", [])
    csv_file_path = state.get("csv_file_path")
    output_dir = state.get("output_dir")
    
    p = Path(csv_file_path)
    pricing_results_path = f"{output_dir}/analytics/analyzed_options.csv"
    sensitivity_test_results_path = f"{output_dir}/analytics/sensitivity_test_results.csv"
    stress_test_results_path = f"{output_dir}/analytics/stress_test_results.csv"
    put_call_parity_path = f"{output_dir}/analytics/put_call_parity.csv"
    final_report_path = state.get("final_report_path")
    
    # Load DataFrames
    df_pricing = pd.read_csv(pricing_results_path)
    sensitivity_test_results = pd.read_csv(sensitivity_test_results_path)
    stress_test_results = pd.read_csv(stress_test_results_path)
    put_call_parity = pd.read_csv(put_call_parity_path)

    ## Set parameters
    report_params = {
        "title": "Ongoing Monitoring Analysis Report",
        "model_name": "Option Pricing, BSM",
        "author_name": "John Doe",
        "group_name": "Front Desk Modeling and Analytics",
        "version": "v1.0"
    }


    llm = get_llm()
    doc = Document()

    # 1. Title Page
    print(">>>>>> [Report Generator Agent] Compiling title page...")
    _add_title_page(doc, report_params)

    # 2. Section 1
    section_ordering = 1
    print(">>>>>> [Report Generator Agent] Compiling section {section_ordering}...")
    _add_section_1(doc, f"{section_ordering}. Introduction", None)

    # 3. Section 2 (Loop per asset)
    section_ordering = 2
    print(f">>>>>> [Report Generator Agent] Compiling section {section_ordering}...")
    doc.add_heading(f"{section_ordering}. Summary of Analysis", level=1)
    
    print(f">>>>>>>>> [Report Generator Agent] Compiling section {section_ordering}: summary table...")
    
    asset = "Equity"
    df_pricing_sub = df_pricing[df_pricing["asset_class"] == asset]
    sensitivity_test_results_sub = sensitivity_test_results[sensitivity_test_results["asset_class"] == asset]
    stress_test_results_sub = stress_test_results[stress_test_results["asset_class"] == asset]
    put_call_parity_sub = put_call_parity[put_call_parity["asset_class_put"] == asset]

    _generate_summary_table(doc, section_ordering = section_ordering, asset=asset)
    doc.add_paragraph("")

    subsection_level = 2
    subsection_ordering = f"{section_ordering}.1"
    print(f">>>>>>>>> [Report Generator Agent] Compiling section {subsection_ordering}: Diagnostic Test...")
    _generate_dignostic_summary(
        doc, 
        llm, 
        asset, 
        df_pricing=df_pricing_sub, 
        df_parity=put_call_parity_sub,
        section_ordering=subsection_ordering,
        section_level = subsection_level
    )

    subsection_level = 2
    subsection_ordering = f"{section_ordering}.2"
    print(f">>>>>>>>> [Report Generator Agent] Compiling section {subsection_ordering}: Sensitivity Test...")
    _generate_sensitivity_test_summary(
        doc, 
        llm, 
        asset, 
        df=sensitivity_test_results_sub, 
        section_ordering=subsection_ordering,
        section_level = subsection_level
    )

    subsection_level = 2
    subsection_ordering = f"{section_ordering}.3"
    print(f">>>>>>>>> [Report Generator Agent] Compiling section {subsection_ordering}: Stress Test...")
    _generate_stress_test_summary(
        doc, 
        llm, 
        asset, 
        df=stress_test_results_sub, 
        section_ordering=subsection_ordering,
        section_level = subsection_level
    )

    doc.save(final_report_path)
    print(f">>> Report saved to {final_report_path}")

    state["errors"] = errors
    state["final_report_path"] = final_report_path
    
    return state

def _add_title_page(doc, params: Dict[str, str]):
    """Add title page and TOC"""
    # Title
    title_run = doc.add_paragraph().add_run(params["title"])
    title_run.bold = True
    title_run.font.size = doc.styles["Title"].font.size
    doc.paragraphs[-1].alignment = 1  # 1=center

    # Subtitle (Model Name)
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(params["model_name"])
    subtitle_para.alignment = 1
    subtitle_run.bold = True

    doc.add_paragraph("")  # spacer

    # Author
    author_para = doc.add_paragraph()
    author_para.add_run("Author: ").bold = True
    author_para.add_run(params["author_name"]).italic = False
    author_para.alignment = 1

    # Group
    group_para = doc.add_paragraph()
    group_para.add_run("Group: ").bold = True
    group_para.add_run(params["group_name"])
    group_para.alignment = 1

    # Report date
    date_str = datetime.date.today().strftime("%B %d, %Y")
    date_para = doc.add_paragraph()
    date_para.add_run("Report Date: ").bold = True
    date_para.add_run(date_str)
    date_para.alignment = 1

    # Version
    version_para = doc.add_paragraph()
    version_para.add_run("Document Version: ").bold = True
    version_para.add_run(params["version"])
    version_para.alignment = 1

    doc.add_page_break()

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    toc_paragraph = doc.add_paragraph()
    toc_field = OxmlElement("w:fldSimple")
    toc_field.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    toc_paragraph._p.append(toc_field)
    doc.add_page_break()


def _add_section_1(doc, heading: str, paragraph: Optional[str]):
    """Add Introduction Section"""
    if not paragraph:
        paragraph = (
            "This section provides contextual background, objectives, and relevant "
            "considerations for the ongoing monitoring analysis. Subsequent sections "
            "expand on methodology, insights, and results."
        )

    doc.add_heading(heading, level=1)
    doc.add_paragraph(paragraph)



def _generate_summary_table(
    doc, 
    section_ordering: int, 
    asset: str
):
    """
    Generates a summary table with 4 columns:
    1. Seq #
    2. Test
    3. Material
    4. Test Conclusion
    """
    paragraph = (
        "In this section, we will provide a summary of tests to evaluate the performance of the BSM model. "
    )
    doc.add_paragraph(paragraph)
    doc.add_paragraph("")
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Seq #'
    hdr_cells[1].text = 'Test'
    hdr_cells[2].text = 'Material'
    hdr_cells[3].text = 'Test Conclusion'
    
    rows_data = [
        ("1", "Diagnostic Test", f"See Section {section_ordering}.1", ""),
        ("2", "Sensitivity Test", f"See Section {section_ordering}.2", ""),
        ("3", "Stress Test", f"See Section {section_ordering}.3", ""),
    ]
    
    for seq, test_name, material, conclusion in rows_data:
        row_cells = table.add_row().cells
        row_cells[0].text = seq
        row_cells[1].text = test_name
        row_cells[2].text = material
        row_cells[3].text = conclusion  # Placeholder

      # Spacer after table



def _generate_dignostic_summary(
    doc, 
    llm, 
    asset: str, 
    df_pricing: pd.DataFrame, 
    df_parity: pd.DataFrame, 
    section_ordering: str,
    section_level: int,
):
    """Generate pricing summary and validation section"""
    doc.add_heading(f"{section_ordering} Summary of Diagnostic Test for {asset}", level=section_level)

    print(f">>>>>>>>>>>> [Report Generator Agent] Compiling section {section_ordering}.1: Data Quality Test...")
    
    _generate_pricing_summary(
        doc,
        llm,
        asset,
        df_pricing,
        section_ordering,
        section_level,
    )

    print(f">>>>>>>>>>>> [Report Generator Agent] Compiling section {section_ordering}.2: Put/Call Parity Test...")
    _generate_parity_summary(
        doc,
        llm,
        asset,
        df_parity,
        section_ordering,
        section_level,
    )


def _generate_pricing_summary(
    doc, 
    llm, 
    asset: str, 
    df: pd.DataFrame, 
    section_ordering: str,
    section_level: int,
):
    df = df.sort_values("T").dropna(subset=['price', 'T'])
    df_str = df.round(4).to_csv(index=False)
    system_prompt_text = (
        "You are an expert Quantitative Analyst writing a formal component for an Ongoing Performance Analysis (OPA) Report. "
        "Your output will be directly embedded into a professional document for senior management.\n\n"
        "### STRICT GUIDELINES:\n"
        "1. **Tone**: Use strictly professional, objective, and formal financial language. (e.g., Use 'The data indicates...' instead of 'I think...').\n"
        "2. **No Conversational Fillers**: Do NOT use phrases like 'Okay', 'Let's see', 'Let me check', 'Wait', or 'Here is the analysis'.\n"
        "3. **No Internal Monologue**: Do NOT output your thinking process. Only output the final analytical conclusions.\n"
        "4. **Direct Start**: Start directly with the bullet points or the section content.\n"
        "5. **Formatting**: Use standard Markdown (bolding for key metrics) suitable for a final report."
    )
    sys_msg = SystemMessage(content=system_prompt_text)

    prompt_quality = (
        f"Dataset for {asset}:\n{df_str}\n\n"
        "### TASK: Data Quality & Integrity Analysis\n"
        "Perform a rigorous audit of the variables listed below. For each variable, evaluate: \n"
        "1) Missingness/Nulls. \n"
        "2) Physical Reasonability (e.g., date logic). \n"
        "3) Financial Logic Consistency (Greeks boundaries).\n\n"
        
        "STRICT GUIDELINES:\n"
        "- Start directly with the analysis. No introductory phrases.\n"
        "- Use the format: **Variable Name**: [Concise assessment].\n\n"
        
        "VARIABLES TO AUDIT:\n"
        "- **date**: Check for missingness and if they align with the current reporting period.\n"
        "- **S & K**: Spot and Strike must be positive; check if they are within reasonable market proximity.\n"
        "- **T**: Time to Maturity must be strictly positive (T > 0).\n"
        "- **r & sigma**: Risk-free rate and Volatility should be non-negative and expressed in decimals.\n"
        "- **Delta**: Must be [0, 1] for Calls and [-1, 0] for Puts.\n"
        "- **Gamma & Vega**: Must be strictly non-negative (>= 0) for long positions.\n"
        "- **Theta**: Should typically be negative (<= 0) for long positions due to time decay.\n"
        "- **Rho**: Must align with option type (usually positive for Calls, negative for Puts).\n\n"
        
        "Output Format: Bullet points starting with the **Bolded Variable Name**."
    )

    batch_inputs = [
        [sys_msg, HumanMessage(content=prompt_quality)],
    ]
    results = llm.batch(batch_inputs)
    res_quality = results[0].content

    full_report = (
        f"{"#"*(section_level+1)} {section_ordering}.1 Data Quality & Integrity\n{res_quality}\n\n"
    )

    add_markdown_to_docx(doc, full_report)

    # Add Figure
    doc.add_paragraph(f"The pricing output of {asset} listed in the below table,")
    for opt_type in ["call", "put"]:
        subset = df[df["option_type"] == opt_type]
        if subset.empty:
            continue
            
        fig, ax = plt.subplots(figsize=(10, 6))
        
        sns.lineplot(
            data=subset, 
            x="T", 
            y="price", 
            hue="underlying",  
            marker="o", 
            ax=ax
        )
        
        ax.set_title(f"{asset} {opt_type.capitalize()} Price vs Time to Maturity")
        ax.set_xlabel("Time to Maturity (T)")
        ax.set_ylabel("Option Price")
        ax.grid(True, alpha=0.3)
        ax.legend(title="Underlying", bbox_to_anchor=(1.05, 1), loc='upper left')
        plt.tight_layout()

        img_stream = BytesIO()
        fig.savefig(img_stream, format="png", dpi=200, bbox_inches="tight")
        img_stream.seek(0)
        
        doc.add_picture(img_stream, width=Inches(5.5))
        
        plt.close(fig)      
        img_stream.close()  
       
def _generate_parity_summary(
    doc, 
    llm, 
    asset: str, 
    df: pd.DataFrame, 
    section_ordering: str,
    section_level: int,
):
    """Generate Gamma Positivity Summary"""
    # doc.add_heading(f"{section_ordering}.2 Summary of Put-Call Parity Testing for {asset}", level=3)
    

    df_str = df[["underlying", "is_parity_valid", "arbitrage_opportunity"]].to_csv(index=False)


    system_prompt_text = (
        "You are an expert Quantitative Analyst writing a formal component for an Ongoing Performance Analysis (OPA) Report. "
        "Your output will be directly embedded into a professional document for senior management.\n\n"
        "### STRICT GUIDELINES:\n"
        "1. **Tone**: Use strictly professional, objective, and formal financial language. (e.g., Use 'The data indicates...' instead of 'I think...').\n"
        "2. **No Conversational Fillers**: Do NOT use phrases like 'Okay', 'Let's see', 'Let me check', 'Wait', or 'Here is the analysis'.\n"
        "3. **No Internal Monologue**: Do NOT output your thinking process. Only output the final analytical conclusions.\n"
        "4. **Direct Start**: Start directly with the bullet points or the section content.\n"
        "5. **Formatting**: Use standard Markdown (bolding for key metrics) suitable for a final report."
    )
    sys_msg = SystemMessage(content=system_prompt_text)

    user_prompt = (
        f"Dataset for {asset}:\n{df_str}\n\n"
        "### TASK: Put Call Parity Check\n"
        "Please provide a structured analysis:\n"
        "STRICT GUIDELINES:\n"
        "- Start directly with the analysis. No introductory phrases.\n"
        "Group by the underlying, count the number of put/call pairs and how many of them pass the put-call parity test.\n"
        "If there are any put-call parity violations, show the arbitrage_opportunity.\n"
        "Output Format: Table with columns: underlying, put_call_pairs_count, fail_count, arbitrage_opportunity."
    )

    batch_inputs = [
        [sys_msg, HumanMessage(content=user_prompt)],
    ]
    results = llm.batch(batch_inputs)
    res_quality = results[0].content


    full_report = (
        f"{"#"*(section_level+1)} {section_ordering}.2 Put/Call Parity Check\n{res_quality}\n\n"
    )
    add_markdown_to_docx(doc, full_report)

    img_reg, img_dev = create_parity_charts(df, asset)
    
    if img_reg:
        doc.add_paragraph("Visual verification of the parity relationship:")
        doc.add_picture(img_reg, width=Inches(5))
        doc.add_paragraph(f"Figure: Parity Regression (LHS vs RHS) for {asset}", style="Caption")
        
    if img_dev:
        doc.add_picture(img_dev, width=Inches(5))
        doc.add_paragraph(f"Figure: Deviation Magnitude across Strikes for {asset}", style="Caption")
    
    add_parity_test_plots(doc, df)

def _generate_sensitivity_test_summary(
    doc, 
    llm, 
    asset: str, 
    df: pd.DataFrame, 
    section_ordering: str,
    section_level: int,
):
    doc.add_heading(f"{section_ordering} Summary of Sensitivity Test for {asset}", level=section_level)

    df_spot = aggregate_spot_sensitivity_data(df)
    df_vol = aggregate_vol_sensitivity_data(df)
    
    spot_str = df_spot.to_markdown(index=False) if df_spot is not None else "No Data"
    vol_str = df_vol.to_markdown(index=False) if df_vol is not None else "No Data"

    system_prompt_text = (
        "You are an expert Quantitative Analyst writing a formal component for an Ongoing Performance Analysis (OPA) Report. "
        "Your output will be directly embedded into a professional document for senior management.\n\n"
        "You are analyzing sensitivity test results."
        "### STRICT GUIDELINES:\n"
        "1. **Tone**: Use strictly professional, objective, and formal financial language. (e.g., Use 'The data indicates...' instead of 'I think...').\n"
        "2. **No Conversational Fillers**: Do NOT use phrases like 'Okay', 'Let's see', 'Let me check', 'Wait', or 'Here is the analysis'.\n"
        "3. **No Internal Monologue**: Do NOT output your thinking process. Only output the final analytical conclusions.\n"
        "4. **Direct Start**: Start directly with the bullet points or the section content.\n"
        "5. **Formatting**: Use standard Markdown (bolding for key metrics) suitable for a final report."
    )
    sys_msg = SystemMessage(content=system_prompt_text)

    prompts = [
        # Spot / Gamma
        f"Portfolio Spot Sensitivity for {asset}:\n{spot_str}\n\n"
        "### Task: Spot & Gamma Analysis\n"
        "1. **Directional Bias**: Does the Delta explain the majority of portfolio gain on Up or Down moves?\n"
        "2. **Convexity**: Is the PnL curve convex (Long Gamma) or concave (Short Gamma)? Check if losses accelerate in one direction.\n",
        
        # Vol / Vega
        f"Portfolio Volatility Sensitivity for {asset}:\n{vol_str}\n\n"
        "### Task: Volatility Analysis\n"
        "1. is the relationship between the pnl and vol bump linear?",
        
        # # Rate / Rho
        # f"Portfolio Rate Sensitivity for {asset}:\n{rate_str}\n\n"
        # "### Task: Interest Rate Analysis\n"
        # "1. **Rho Exposure**: How does the portfolio react to rate hikes?\n"
        # "2. **Implication**: Briefly explain what this implies about the positioning (e.g., Net Long/Short)."
    ]

    inputs = [[sys_msg, HumanMessage(content=p)] for p in prompts]
    results = llm.batch(inputs)

    # Run LLM (Batch)
    inputs = [[sys_msg, HumanMessage(content=p)] for p in prompts]
    results = llm.batch(inputs)

    res_spot = results[0].content
    res_vol = results[1].content

    report_spot = f"{"#"*(section_level+1)} {section_ordering}.1 Spot & Gamma Analysis\n{res_spot}\n\n"
    add_markdown_to_docx(doc, report_spot)
    img_spot = create_spot_sensitivity_chart(df_spot)
    if img_spot:
        doc.add_picture(img_spot, width=Inches(5))
        doc.add_paragraph(f"Figure: Spot Price Sensitivity for {asset}", style="Caption")


    report_vol = f"{"#"*(section_level+1)} {section_ordering}.2 Volatility Analysis\n{res_vol}\n\n"
    add_markdown_to_docx(doc, report_vol)
    img_vol = create_vol_sensitivity_chart(df_vol)
    if img_vol:
        doc.add_picture(img_vol, width=Inches(5))
        doc.add_paragraph(f"Figure: Volatility Price Sensitivity for {asset}", style="Caption")
    
    add_comparative_sensitivity_plots(doc, df)

    


def _generate_stress_test_summary(doc, 
    llm, 
    asset: str, 
    df: pd.DataFrame, 
    section_ordering: str,
    section_level: int,
):
    doc.add_heading(f"{section_ordering} Summary of Stress Testing for {asset}", level=section_level)
    
    df_agg_underlying_level, df_agg_scenario_level = prepare_stress_test_data(df)
    df_str = df_agg_scenario_level.to_markdown(index=False)

    system_prompt_text = (
        "You are an expert Quantitative Analyst writing a formal component for an Ongoing Performance Analysis (OPA) Report. "
        "Your output will be directly embedded into a professional document for senior management.\n\n"
        "You are analyzing stress test results."
        "### STRICT GUIDELINES:\n"
        "1. **Tone**: Use strictly professional, objective, and formal financial language. (e.g., Use 'The data indicates...' instead of 'I think...').\n"
        "2. **No Conversational Fillers**: Do NOT use phrases like 'Okay', 'Let's see', 'Let me check', 'Wait', or 'Here is the analysis'.\n"
        "3. **No Internal Monologue**: Do NOT output your thinking process. Only output the final analytical conclusions.\n"
        "4. **Direct Start**: Start directly with the bullet points or the section content.\n"
        "5. **Formatting**: Use standard Markdown (bolding for key metrics) suitable for a final report."
    )
    sys_msg = SystemMessage(content=system_prompt_text)

    user_prompt = (
        f"Stress Test Results for {asset}:\n\n"
        f"{df_str}\n\n"
        "Please provide a Risk Assessment covering:\n"
        "1. Rank the severity of stress scenario and provide the potential reason."
    )
    batch_inputs = [
        [sys_msg, HumanMessage(content=user_prompt)],
    ]
    results = llm.batch(batch_inputs)
    res_quality = results[0].content
    add_markdown_to_docx(doc, res_quality)

    img_stress = create_stress_chart(df_agg_underlying_level)
    if img_stress:
        doc.add_picture(img_stress, width=Inches(5.5))
        doc.add_paragraph(f"Figure: Stress Scenario Impact on {asset}", style="Caption")

    add_cross_asset_stress_plots(doc, df)

    





def _process_inline_formatting(paragraph, text):
    pattern = r'(\*\*.*?\*\*|\*.*?\*)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        
        run = paragraph.add_run()
        
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            clean_text = part[2:-2]
            run.text = clean_text
            run.bold = True
            
        elif part.startswith('*') and part.endswith('*') and len(part) >= 2:
            clean_text = part[1:-1]
            run.text = clean_text
            run.italic = True
            
        else:
            run.text = part

def _create_word_table(doc, table_buffer):
    if not table_buffer:
        return

    raw_data = []
    for row in table_buffer:
        cells = [c.strip() for c in row.split('|')]
        if cells[0] == "": cells.pop(0)
        if cells and cells[-1] == "": cells.pop(-1)
        raw_data.append(cells)

    if not raw_data:
        return

    rows = len(raw_data)
    cols = len(raw_data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'  

    for i, row_data in enumerate(raw_data):
        if i == 1 and all(re.match(r'^[\s\-\:]+$', c) for c in row_data):
            continue 
        word_row_idx = i if i < 1 else (i - 1 if any(re.match(r'^[\s\-\:]+$', c) for c in raw_data[1]) else i)
        
        if word_row_idx >= len(table.rows):
            table.add_row()

        row_cells = table.rows[word_row_idx].cells
        for j, val in enumerate(row_data):
            if j < len(row_cells):
                row_cells[j].paragraphs[0].text = ""
                _process_inline_formatting(row_cells[j].paragraphs[0], val)

def add_markdown_to_docx(doc, text):
    lines = text.split('\n')
    table_buffer = []
    in_table = False
    
    for line in lines:
        line = line.strip()
        
        if line.startswith('|') and line.endswith('|'):
            table_buffer.append(line)
            in_table = True
            continue
        else:
            if in_table:
                _create_word_table(doc, table_buffer)
                table_buffer = []
                in_table = False

        if not line:
            continue
        
        header_index = False
        for header_pattern in ["#"*i for i in range(11,2,-1)]:
            if line.startswith(header_pattern):
                clean_text = line.replace('#', '').strip()
                clean_text = clean_text.replace('**', '').replace('*', '')
                doc.add_heading(clean_text, level=min(len(header_pattern), 9))
                header_index = True
                break
        
        if header_index:
            continue  
        elif line.startswith('**') and line.endswith('**') and len(line) < 80:
            clean_text = line[2:-2].strip()
            doc.add_heading(clean_text, level=3)
        elif line.startswith(('-', '*')) and len(line) > 1 and line[1] == ' ':
            clean_text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            _process_inline_formatting(p, clean_text)
            
        elif re.match(r'^\d+\.\s', line):
            # clean_text = re.sub(r'^\d+\.\s', '', line).strip()
            # p = doc.add_paragraph(style='List Number')
            # _process_inline_formatting(p, clean_text)
            
            # Use List Paragraph to avoid auto-numbering continuity issues
            p = doc.add_paragraph(style='List Paragraph')
            _process_inline_formatting(p, line)
            
        elif line == '---':
            p = doc.add_paragraph()
            p.add_run().add_break()
            
        else:
            p = doc.add_paragraph()
            _process_inline_formatting(p, line)
            
    if in_table:
        _create_word_table(doc, table_buffer)

def aggregate_spot_sensitivity_data(df):
    prefix = f"sensitivity_scen_spot_bump_"
    all_bump_cols = [c for c in df.columns if prefix in c]
    bump_zero_col = f'sensitivity_scen_spot_bump_0'
    if bump_zero_col in df.columns:
        df['base_value'] = df[bump_zero_col]
        analysis_cols = [c for c in all_bump_cols if c != bump_zero_col]
    else:
        df['base_value'] = df['price']
        analysis_cols = all_bump_cols

    id_vars = ['underlying', 'delta', 'gamma', 'S', 'base_value']
    df_long = df.melt(
        id_vars=id_vars, 
        value_vars=analysis_cols,
        var_name='scenario_name', 
        value_name='scenario_price'
    )
    df_long['spot_bump'] = df_long['scenario_name'].str.replace(f'sensitivity_scen_spot_bump_', '').astype(float)
    df_long['dS'] = df_long['S'] * df_long['spot_bump']
    df_long['real_pnl'] = df_long['scenario_price'] - df_long['base_value']
    df_long['delta_pnl'] = df_long['delta'] * df_long['dS']
    df_long['gamma_pnl'] = 0.5 * df_long['gamma'] * (df_long['dS']**2)
    df_long['expect_pnl'] = df_long['delta_pnl'] + df_long['gamma_pnl']
    df_long['residual_pnl'] = df_long['real_pnl'] - df_long['expect_pnl']
    final_cols = ['underlying', 'spot_bump', 'real_pnl', 'delta_pnl', 'gamma_pnl', 'expect_pnl', 'residual_pnl']
    df_agg = df_long[final_cols].groupby(['spot_bump'])[[
        'real_pnl', 
        'delta_pnl', 
        'gamma_pnl', 
        'expect_pnl'
    ]].sum().reset_index()
    return df_agg

def create_spot_sensitivity_chart(df_agg):
    df_plot = df_agg.sort_values('spot_bump')

    bumps = [f"{int(b*100)}%" for b in df_plot['spot_bump']]
    delta_pnl = df_plot['delta_pnl']
    gamma_pnl = df_plot['gamma_pnl']
    real_pnl = df_plot['real_pnl']

    plt.figure(figsize=(7, 4.5))
    
    p1 = plt.bar(bumps, delta_pnl, color='#1f77b4', alpha=0.7, label='Delta PnL')
    p2 = plt.bar(bumps, gamma_pnl, bottom=delta_pnl, color='#ff7f0e', alpha=0.7, label='Gamma PnL')
    
    plt.scatter(bumps, real_pnl, color='black', marker='D', s=40, zorder=3, label='Real PnL (Observed)')

    plt.axhline(0, color='black', linewidth=0.8, linestyle='-')
    
    plt.title(f"PnL Attribution Breakdown:", fontsize=12, weight='bold')
    plt.xlabel("Spot Price Bump (%)", fontsize=10)
    plt.ylabel("Profit & Loss ($)", fontsize=10)
    plt.legend(loc='best', fontsize=9)
    plt.grid(axis='y', linestyle='--', alpha=0.4)

    plt.tight_layout()

    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150)
    plt.close()
    img_buffer.seek(0)
    
    return img_buffer

def aggregate_vol_sensitivity_data(df):
    prefix = f"sensitivity_scen_vol_bump_"
    all_bump_cols = [c for c in df.columns if prefix in c]
    bump_zero_col = f'sensitivity_scen_vol_bump_0'
    if bump_zero_col in df.columns:
        df['base_value'] = df[bump_zero_col]
        analysis_cols = [c for c in all_bump_cols if c != bump_zero_col]
    else:
        df['base_value'] = df['price']
        analysis_cols = all_bump_cols
    id_vars = ['underlying', 'base_value']
    df_long = df.melt(
        id_vars=id_vars, 
        value_vars=analysis_cols,
        var_name='scenario_name', 
        value_name='scenario_price'
    )
    df_long['vol_bump'] = df_long['scenario_name'].str.replace(f'sensitivity_scen_vol_bump_', '').astype(float)
    df_long['real_pnl'] = df_long['scenario_price'] - df_long['base_value']
    final_cols = ['underlying', 'vol_bump', 'real_pnl']
    df_agg = df_long[final_cols].groupby(['vol_bump'])[[
        'real_pnl', 
    ]].sum().reset_index()
    return df_agg

def create_vol_sensitivity_chart(df_agg):
    
    plt.figure(figsize=(6, 3.5)) # Compact size for Word doc

    plt.plot(df_agg['vol_bump'], df_agg['real_pnl'], marker='o', linewidth=2, color='#1f77b4', label='PnL')
    plt.axhline(0, color='black', linewidth=0.8, linestyle='--')
    plt.axvline(0, color='black', linewidth=0.8, linestyle='--')
    plt.fill_between(df_agg['vol_bump'], df_agg['real_pnl'], 0, where=(df_agg['real_pnl']>=0), 
                        facecolor='green', alpha=0.1, interpolate=True)
    plt.fill_between(df_agg['vol_bump'], df_agg['real_pnl'], 0, where=(df_agg['real_pnl']<0), 
                        facecolor='red', alpha=0.1, interpolate=True)
    scenario_label = 'Vol'
    plt.title(f"{scenario_label} Sensitivity PnL", fontsize=10, weight='bold')
    plt.xlabel(f"{scenario_label} Bump", fontsize=9)
    plt.ylabel("Portfolio PnL ($)", fontsize=9)
    plt.grid(True, alpha=0.3, linestyle=':')
    plt.xticks(fontsize=8)
    plt.yticks(fontsize=8)
    plt.tight_layout()

    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150)
    plt.close()
    img_buffer.seek(0)
    
    return img_buffer




def prepare_stress_test_data(df):
    prefix = "stress_scen_"
    cols = [c for c in df.columns if prefix in c]
    
    id_vars = ['underlying','price']
    df_long = df.melt(
        id_vars=id_vars, 
        value_vars=cols,
        var_name='scenario_name', 
        value_name='scenario_price'
    )
    df_long['scenario_name'] = df_long['scenario_name'].str.replace(f'stress_scen_', '')
    df_long['PnL'] = df_long['scenario_price'] - df_long['price']
    df_agg_underlying_level = df_long.groupby(['underlying','scenario_name'])[[
        'price',
        'PnL', 
    ]].sum().reset_index()
    df_agg_underlying_level['PnL %'] = (df_agg_underlying_level['PnL'] / df_agg_underlying_level['price']*100).round(2).astype(str) + "%"
    df_agg_scenario_level = df_long.groupby(['scenario_name'])[[
        'price',
        'PnL', 
    ]].sum().reset_index()
    df_agg_scenario_level['PnL %'] = (df_agg_scenario_level['PnL'] / df_agg_scenario_level['price']*100).round(2).astype(str) + "%"
    return df_agg_underlying_level, df_agg_scenario_level

def create_stress_chart(df_agg):
    """
    Generates a Horizontal Bar Chart for Stress Scenarios.
    Red bars for losses, Green for gains.
    """
    if df_agg is None or df_agg.empty:
        return None
        
    plt.figure(figsize=(10, 6))
    sns.barplot(data=df_agg, x='scenario_name', y='PnL', hue='underlying', palette='viridis')

    plt.axhline(0, color='black', linewidth=0.8) 
    plt.title("Portfolio Stress Test: PnL by Scenario & Underlying", weight='bold')
    plt.xticks(rotation=45)
    plt.ylabel("PnL ($)")
    plt.grid(axis='y', linestyle='--', alpha=0.5)
    plt.tight_layout()
    
    # Save to buffer
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150)
    plt.close()
    img_buffer.seek(0)
    
    return img_buffer

def add_comparative_sensitivity_plots(doc, df_sensitivity):
    factors = {
        'spot': 'Spot Price Bump (%)',
        'vol': 'Volatility Bump (abs)',
        'rate': 'Interest Rate Bump (abs)'
    }
    
    for opt_type in ["call", "put"]:
        subset_orig = df_sensitivity[df_sensitivity["option_type"] == opt_type]
        if subset_orig.empty:
            continue

        fig, axes = plt.subplots(1, 3, figsize=(20, 6))
        fig.suptitle(f"Cross-Asset {opt_type.upper()} Sensitivity Comparison", fontsize=16)

        for i, (factor, label) in enumerate(factors.items()):
            scen_cols = [c for c in df_sensitivity.columns if f"scen_{factor}_bump" in c]
            
            comparison_data = []
            for _, row in subset_orig.iterrows():
                for col in scen_cols:
                    try:
                        bump_val = float(col.split('_bump_')[-1])
                        comparison_data.append({
                            'Bump': bump_val,
                            'Price': row[col],
                            'Underlying': row['underlying'],
                            'ID': row['ID']
                        })
                    except: continue
            
            df_plot = pd.DataFrame(comparison_data).sort_values('Bump')

            sns.lineplot(
                data=df_plot, 
                x='Bump', 
                y='Price', 
                hue='Underlying', 
                marker='o', 
                ax=axes[i]
            )
            
            axes[i].set_title(f"Sensitivity to {factor.capitalize()}")
            axes[i].set_xlabel(label)
            axes[i].set_ylabel("Option Price")
            axes[i].grid(True, alpha=0.3)
            axes[i].axvline(x=0, color='black', linestyle='--', alpha=0.5)

        plt.tight_layout(rect=[0, 0.03, 1, 0.95])

        img_stream = BytesIO()
        fig.savefig(img_stream, format="png", dpi=200, bbox_inches="tight")
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(6.2))
        
        plt.close(fig)
        img_stream.close()


def add_cross_asset_stress_plots(doc, df_stress):
    stress_cols = [c for c in df_stress.columns if c.startswith('stress_scen_')]
    
    for opt_type in ["call", "put"]:
        subset = df_stress[df_stress["option_type"] == opt_type]
        if subset.empty:
            continue
            
        plot_data = []
        for _, row in subset.iterrows():
            plot_data.append({
                'Underlying': row['underlying'],
                'Price': row['price'],
                'Scenario': 'Baseline'
            })
            for col in stress_cols:
                scen_name = col.replace('stress_scen_', '')
                plot_data.append({
                    'Underlying': row['underlying'],
                    'Price': row[col],
                    'Scenario': scen_name
                })
        
        df_plot = pd.DataFrame(plot_data)

        plt.figure(figsize=(14, 8))
        ax = sns.barplot(
            data=df_plot, 
            x='Underlying', 
            y='Price', 
            hue='Scenario',
            palette='tab10'
        )
        
        plt.title(f"Cross-Asset Impact Analysis: {opt_type.upper()} Options", fontsize=15)
        plt.ylabel("Option Price")
        plt.xlabel("Underlying Asset")
        plt.grid(axis='y', linestyle='--', alpha=0.5)
        
        plt.legend(title="Stress Scenarios", bbox_to_anchor=(1.05, 1), loc='upper left')
        plt.tight_layout()

        img_stream = BytesIO()
        plt.savefig(img_stream, format="png", dpi=200, bbox_inches="tight")
        img_stream.seek(0)
        
        doc.add_heading(f"{opt_type.capitalize()} Stress Exposure by Underlying", level=2)
        doc.add_picture(img_stream, width=Inches(6.2))
        
        plt.close()
        img_stream.close()

def prepare_parity_summary(df):
    """
    Calculates statistical summaries of Put-Call Parity checks.
    Returns a dictionary of stats and a Markdown string for the LLM.
    """
    total_pairs = len(df)
    if total_pairs == 0:
        return None, "No data available."
    
    # 1. Basic Stats
    num_valid = df['is_parity_valid'].sum()
    pass_rate = (num_valid / total_pairs) * 100
    
    mean_diff = df['abs_diff'].mean()
    max_diff = df['abs_diff'].max()
    median_diff = df['abs_diff'].median()
    
    # 2. Arbitrage Opportunities (Assume threshold > 0.1 or specific column logic)
    # If 'arbitrage_opportunity' column is boolean:
    arb_ops = df[df['arbitrage_opportunity'] == True]
    num_arbs = len(arb_ops)
    
    # 3. Identify Top Violations (Worst 5)
    # Sort by absolute difference descending
    worst_cases = df.sort_values('abs_diff', ascending=False).head(5)
    
    # Format worst cases for LLM
    # We select key columns to keep tokens low
    cols_to_show = ['K', 'T', 'price_call', 'price_put', 'abs_diff', 'is_parity_valid']
    worst_table = worst_cases[cols_to_show].round(4).to_markdown(index=False)
    
    # 4. Construct Summary String
    summary_md = (
        f"### Parity Statistics\n"
        f"- **Total Pairs Analyzed**: {total_pairs}\n"
        f"- **Pass Rate**: {pass_rate:.2f}% ({num_valid}/{total_pairs})\n"
        f"- **Mean Deviation**: {mean_diff:.4f}\n"
        f"- **Max Deviation**: {max_diff:.4f}\n"
        f"- **Flagged Arbitrage Opps**: {num_arbs}\n\n"
        f"### Top 5 Worst Violations (Largest Diff)\n"
        f"{worst_table}"
    )
    
    return df, summary_md



def create_parity_charts(df, asset_name):
    """
    Generates 2 charts: 
    1. LHS vs RHS Scatter (The Parity Line)
    2. Strike vs Deviation (Error Distribution)
    """
    if df is None or df.empty:
        return None, None

    # Chart 1: Parity Regression (LHS vs RHS)
    plt.figure(figsize=(5, 4))
    
    # Plot perfect parity line
    max_val = max(df['lhs'].max(), df['rhs'].max())
    min_val = min(df['lhs'].min(), df['rhs'].min())
    plt.plot([min_val, max_val], [min_val, max_val], color='black', linestyle='--', alpha=0.5, label='Theoretical Parity')
    
    # Scatter plot of actual data
    # Color by validity (Green=Good, Red=Bad)
    colors = df['is_parity_valid'].map({True: 'green', False: 'red'})
    plt.scatter(df['lhs'], df['rhs'], c=colors, s=15, alpha=0.6, label='Observed Pairs')
    
    plt.title(f"{asset_name}: Put-Call Parity Regression", fontsize=10)
    plt.xlabel("LHS (Call + PV(K))")
    plt.ylabel("RHS (Put + Spot)")
    plt.grid(True, alpha=0.3)
    plt.legend()
    plt.tight_layout()
    
    img_regression = io.BytesIO()
    plt.savefig(img_regression, format='png', dpi=150)
    plt.close()
    img_regression.seek(0)
    
    # Chart 2: Deviation by Strike
    plt.figure(figsize=(5, 4))
    plt.scatter(df['K'], df['abs_diff'], color='purple', s=15, alpha=0.6)
    plt.axhline(0, color='black', linewidth=0.8)
    
    plt.title(f"{asset_name}: Parity Deviation by Strike", fontsize=10)
    plt.xlabel("Strike Price (K)")
    plt.ylabel("Absolute Difference ($)")
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    
    img_deviation = io.BytesIO()
    plt.savefig(img_deviation, format='png', dpi=150)
    plt.close()
    img_deviation.seek(0)
    
    return img_regression, img_deviation


def add_parity_test_plots(doc, df_parity):
    sns.set_theme(style="whitegrid")
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 6))

    sns.scatterplot(data=df_parity, x='lhs', y='rhs', hue='underlying', s=100, ax=ax1)
    
    max_val = max(df_parity['lhs'].max(), df_parity['rhs'].max())
    min_val = min(df_parity['lhs'].min(), df_parity['rhs'].min())
    ax1.plot([min_val, max_val], [min_val, max_val], color='red', linestyle='--', alpha=0.6, label='Ideal Parity')
    
    ax1.set_title("Put-Call Parity: LHS vs RHS", fontsize=14)
    ax1.set_xlabel("LHS (Call + PV of Strike)")
    ax1.set_ylabel("RHS (Spot + Put)")
    ax1.legend()

    sns.barplot(data=df_parity, x='ID_call', y='abs_diff', hue='underlying', ax=ax2)
    ax2.set_title("Arbitrage Deviation (Abs Diff)", fontsize=14)
    ax2.set_ylabel("Difference Value")
    ax2.set_xlabel("Pair ID")
    
    plt.tight_layout()

    img_stream = BytesIO()
    fig.savefig(img_stream, format="png", dpi=200, bbox_inches="tight")
    img_stream.seek(0)
    
    doc.add_heading("Put-Call Parity Validation Analysis", level=2)
    doc.add_picture(img_stream, width=Inches(6.0))
    
    plt.close(fig)
    img_stream.close()