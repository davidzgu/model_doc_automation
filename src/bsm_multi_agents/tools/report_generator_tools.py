# -*- coding: utf-8 -*-
"""
Word report generation tools for creating professional OPA documents.

Generates Word documents with markdown content and embedded charts.
"""
import json
from pathlib import Path
from typing import Union, List, Dict, Any
from datetime import datetime
from langchain_core.tools import tool
from .tool_registry import register_tool

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import markdown
    from bs4 import BeautifulSoup
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@register_tool(tags=["report", "word", "document"], roles=["report_generator"])
@tool("create_word_report")
def create_word_report(
    markdown_path: str,
    chart_paths: Union[str, List[str]],
    output_path: str = "output/report.docx",
    title: str = "Ongoing Performance Analysis Report"
) -> str:
    """
    Create a professional Word document from markdown summary and charts.

    Args:
        markdown_path: Path to the markdown summary file
        chart_paths: Path(s) to chart images (JSON string of list or single path)
        output_path: Path where the Word document will be saved
        title: Document title

    Returns:
        JSON string with document path and status
    """
    if not DOCX_AVAILABLE:
        return json.dumps({
            "status": "error",
            "message": "python-docx not available. Install with: pip install python-docx markdown beautifulsoup4"
        })

    try:
        # Parse chart_paths if it's a JSON string
        if isinstance(chart_paths, str):
            try:
                chart_paths = json.loads(chart_paths)
            except json.JSONDecodeError:
                chart_paths = [chart_paths]

        if not isinstance(chart_paths, list):
            chart_paths = [chart_paths]

        # Read markdown content
        md_path = Path(markdown_path)
        if not md_path.exists():
            return json.dumps({
                "status": "error",
                "message": f"Markdown file not found: {markdown_path}"
            })

        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Create Word document
        doc = Document()

        # Set document margins (narrower for more content)
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Set up professional styles
        _setup_document_styles(doc)

        # Add cover page
        _add_cover_page(doc, title)

        # Add table of contents placeholder
        doc.add_page_break()
        _add_toc_placeholder(doc)

        # Add executive summary section
        doc.add_page_break()

        # Parse and add markdown content
        _add_markdown_content(doc, md_content)

        # Add charts section
        if chart_paths:
            doc.add_page_break()

            # Charts section header with decorative line
            charts_heading = doc.add_heading('Visualization Charts', level=1)
            _add_horizontal_line(doc)
            doc.add_paragraph()

            for i, chart_path in enumerate(chart_paths, 1):
                chart_file = Path(chart_path)
                if chart_file.exists():
                    # Add chart number and caption
                    caption = chart_file.stem.replace('_', ' ').title()
                    chart_heading = doc.add_heading(f"Chart {i}: {caption}", level=2)

                    # Add image with border
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    picture = run.add_picture(str(chart_file), width=Inches(6.0))

                    # Add image caption
                    caption_para = doc.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.add_run(f"Figure {i}: {caption}")
                    caption_run.font.size = Pt(9)
                    caption_run.font.italic = True
                    caption_run.font.color.rgb = RGBColor(128, 128, 128)

                    doc.add_paragraph()  # Spacing
                else:
                    doc.add_paragraph(f"⚠️ Chart not found: {chart_path}")

        # Add footer with page numbers
        _add_page_numbers(doc)

        # Ensure output directory exists
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        # Save document
        doc.save(str(output_file))

        return json.dumps({
            "status": "success",
            "document_path": str(output_file),
            "message": f"Professional Word report generated successfully"
        })

    except Exception as e:
        return json.dumps({
            "status": "error",
            "message": f"Error creating Word report: {str(e)}"
        })


def _add_cover_page(doc, title):
    """Add a professional cover page."""
    # Title (centered, large, bold)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue

    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Quantitative Analysis Report")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)

    # Add spacing
    for _ in range(3):
        doc.add_paragraph()

    # Date and time
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated on:\n{datetime.now().strftime('%B %d, %Y')}\n{datetime.now().strftime('%H:%M:%S')}")
    date_run.font.size = Pt(12)
    date_run.font.italic = True
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    # Add more spacing
    for _ in range(8):
        doc.add_paragraph()

    # Footer info
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("CONFIDENTIAL\nFor Internal Use Only")
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)


def _add_toc_placeholder(doc):
    """Add table of contents placeholder."""
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Note: Actual TOC generation requires field codes which is complex
    # Add placeholder text
    toc_para = doc.add_paragraph()
    toc_run = toc_para.add_run(
        "1. Executive Summary\n"
        "2. Portfolio Overview\n"
        "3. Validation Results\n"
        "4. Greeks Analysis\n"
        "5. Visualization Charts\n"
        "6. Recommendations"
    )
    toc_run.font.size = Pt(11)
    toc_run.font.color.rgb = RGBColor(80, 80, 80)


def _add_horizontal_line(doc):
    """Add a decorative horizontal line."""
    para = doc.add_paragraph()
    run = para.add_run()
    # Add bottom border to simulate horizontal line
    p = para._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # Line thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '003366')  # Dark blue
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_page_numbers(doc):
    """Add page numbers to footer."""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page number field
    run = footer_para.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Add page number text
    footer_para.text = "Page "
    # Note: Actual page number field requires complex XML manipulation
    # For simplicity, we just add placeholder text


def _setup_document_styles(doc):
    """Set up custom styles for the document."""
    styles = doc.styles

    # Heading 1 style - Main sections
    if 'Heading 1' in styles:
        h1 = styles['Heading 1']
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(6)

    # Heading 2 style - Subsections
    if 'Heading 2' in styles:
        h2 = styles['Heading 2']
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 102, 204)  # Medium blue
        h2.paragraph_format.space_before = Pt(10)
        h2.paragraph_format.space_after = Pt(4)

    # Heading 3 style - Sub-subsections
    if 'Heading 3' in styles:
        h3 = styles['Heading 3']
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(51, 102, 153)
        h3.paragraph_format.space_before = Pt(8)
        h3.paragraph_format.space_after = Pt(3)

    # Normal text style
    if 'Normal' in styles:
        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        normal.paragraph_format.space_after = Pt(6)


def _add_markdown_content(doc, md_content):
    """Parse markdown and add to Word document with enhanced formatting."""
    # Convert markdown to HTML
    html = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'nl2br'])
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup.children:
        if element.name == 'h1':
            heading = doc.add_heading(element.get_text(), level=1)
            # Add decorative line under H1
            _add_horizontal_line(doc)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            # Check if paragraph contains special markers (✅, ❌, ⚠️)
            text = element.get_text()
            para = doc.add_paragraph()

            if '✅' in text or '❌' in text or '⚠️' in text:
                # Highlight important status messages
                run = para.add_run(text)
                if '✅' in text:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                elif '❌' in text:
                    run.font.color.rgb = RGBColor(192, 0, 0)  # Red
                elif '⚠️' in text:
                    run.font.color.rgb = RGBColor(255, 140, 0)  # Orange
                run.font.bold = True
            else:
                para.add_run(text)

        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                para = doc.add_paragraph(li.get_text(), style='List Bullet')
                # Highlight bullets with special markers
                text = li.get_text()
                if '**' in str(li):  # Bold text
                    for run in para.runs:
                        if text.startswith('**'):
                            run.font.bold = True

        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='List Number')

        elif element.name == 'table':
            _add_styled_table(doc, element)

        elif element.name == 'pre':
            # Code block with background
            code_para = doc.add_paragraph(element.get_text())
            code_para.style = 'Normal'
            for run in code_para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue
            # Add light gray background (approximation)
            code_para.paragraph_format.left_indent = Inches(0.25)
            code_para.paragraph_format.right_indent = Inches(0.25)


def _add_styled_table(doc, table_element):
    """Add an HTML table with professional styling to the Word document."""
    rows = table_element.find_all('tr')
    if not rows:
        return

    # Get dimensions
    num_rows = len(rows)
    num_cols = len(rows[0].find_all(['th', 'td']))

    # Create table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Medium Shading 1 Accent 1'  # Professional blue theme

    # Set column widths for better readability
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(6.5 / num_cols)

    # Fill table with enhanced styling
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            table_cell = table.rows[i].cells[j]
            table_cell.text = cell.get_text().strip()

            # Style header row
            if i == 0:
                for paragraph in table_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                # Set header background color (dark blue)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '003366')
                table_cell._element.get_or_add_tcPr().append(shading_elm)
            else:
                # Alternate row colors for better readability
                for paragraph in table_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

                if i % 2 == 0:
                    # Light blue for even rows
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'E8F2F7')
                    table_cell._element.get_or_add_tcPr().append(shading_elm)

    # Add some spacing after table
    doc.add_paragraph()


def _add_table(doc, table_element):
    """Legacy table function - redirects to styled version."""
    _add_styled_table(doc, table_element)
